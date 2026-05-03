import os
import time
import json
import uuid
import urllib.request
import http.client
import sqlite3

from reportlab.pdfgen import canvas
from docx import Document

def generate_pdf(filename):
    c = canvas.Canvas(filename)
    text = "This is a test document. " * 30 # > 100 words
    c.drawString(100, 750, "Page 1 Content")
    c.drawString(100, 700, text)
    c.showPage()
    
    text2 = "Here is another page with more information about the processing pipeline. " * 30
    c.drawString(100, 750, "Page 2 Content")
    c.drawString(100, 700, text2)
    c.showPage()
    
    c.save()
    print(f"Generated {filename}")

def generate_docx(filename):
    doc = Document()
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph('This is the introductory section of the document. ' * 20)
    
    doc.add_heading('Methodology', level=1)
    doc.add_paragraph('This section describes the methods used in our advanced research. ' * 20)
    
    doc.save(filename)
    print(f"Generated {filename}")

def upload_and_check(filepath, kb_id):
    filename = os.path.basename(filepath)
    boundary = '----TestBoundary123'
    
    with open(filepath, 'rb') as f:
        file_data = f.read()
        
    ext = os.path.splitext(filename)[1].lower()
    content_type = 'application/pdf' if ext == '.pdf' else 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        
    body = (
        f'--{boundary}\r\n'
        f'Content-Disposition: form-data; name="file"; filename="{filename}"\r\n'
        f'Content-Type: {content_type}\r\n\r\n'.encode('utf-8') +
        file_data +
        f'\r\n--{boundary}--\r\n'.encode('utf-8')
    )

    conn = http.client.HTTPConnection('127.0.0.1', 8000)
    conn.request(
        'POST',
        f'/api/v1/ingest/file?kb_id={kb_id}',
        body=body,
        headers={'Content-Type': f'multipart/form-data; boundary={boundary}'}
    )
    resp = conn.getresponse()
    data = json.loads(resp.read().decode())
    print(f"\nPOST {filename} => {resp.status}")
    print(json.dumps(data, indent=2))
    conn.close()
    
    if resp.status == 200:
        job_id = data['job_id']
        
        # Wait up to 10 seconds for completion
        for _ in range(10):
            time.sleep(2)
            try:
                resp2 = urllib.request.urlopen(f'http://127.0.0.1:8000/api/v1/ingest/jobs/{job_id}')
                job_data = json.loads(resp2.read().decode())
                if job_data['status'] in ['complete', 'failed']:
                    print(f"\nGET /jobs/{job_id} =>")
                    print(json.dumps(job_data, indent=2))
                    break
                else:
                    print(f"Status: {job_data['status']}... waiting")
            except Exception as e:
                print(f"Error checking job: {e}")
                
        return data['document_id']
    return None

if __name__ == '__main__':
    pdf_file = "test_doc.pdf"
    docx_file = "test_doc.docx"
    
    generate_pdf(pdf_file)
    generate_docx(docx_file)
    
    # Create test user and KB
    test_user_id = str(uuid.uuid4())
    test_kb_id = str(uuid.uuid4())
    
    db = sqlite3.connect('nexus.db')
    db.execute(
        "INSERT INTO users (id, email, hashed_password, role, is_active, created_at, updated_at) VALUES (?, ?, ?, ?, ?, datetime('now'), datetime('now'))",
        (test_user_id, f'test_{test_user_id}@nexus.dev', 'hashed_pw', 'admin', 1)
    )
    db.execute(
        "INSERT INTO knowledge_bases (id, name, description, owner_id, created_at, updated_at) VALUES (?, ?, ?, ?, datetime('now'), datetime('now'))",
        (test_kb_id, 'Test KB', 'For testing ingestion', test_user_id)
    )
    db.commit()
    
    print("\n--- Uploading PDF ---")
    doc_id_pdf = upload_and_check(pdf_file, test_kb_id)
    
    print("\n--- Uploading DOCX ---")
    doc_id_docx = upload_and_check(docx_file, test_kb_id)
    
    print("\n--- Testing Query Endpoint (Text) ---")
    query_body = json.dumps({
        "kb_id": test_kb_id,
        "query": "What are the methods used in research?",
        "top_k": 3,
        "modality_filter": "text"
    })
    conn = http.client.HTTPConnection('127.0.0.1', 8000)
    conn.request('POST', '/api/v1/query/', body=query_body, headers={'Content-Type': 'application/json'})
    resp3 = conn.getresponse()
    query_data = json.loads(resp3.read().decode())
    print(f"POST /query => {resp3.status}")
    print(json.dumps(query_data, indent=2))
    conn.close()
    
    # Generate test image
    image_file = "test_image.png"
    from PIL import Image, ImageDraw, ImageFont
    img = Image.new('RGB', (400, 200), color = (73, 109, 137))
    d = ImageDraw.Draw(img)
    d.text((10,10), "NEXUS Visual Search Architecture", fill=(255,255,0))
    img.save(image_file)
    
    print("\n--- Uploading PNG Image ---")
    doc_id_image = upload_and_check(image_file, test_kb_id)
    
    print("\n--- Testing Query Endpoint (Image) ---")
    query_body_img = json.dumps({
        "kb_id": test_kb_id,
        "query": "A diagram about system architecture",
        "top_k": 3,
        "modality_filter": "image"
    })
    conn = http.client.HTTPConnection('127.0.0.1', 8000)
    conn.request('POST', '/api/v1/query/', body=query_body_img, headers={'Content-Type': 'application/json'})
    resp_img = conn.getresponse()
    query_data_img = json.loads(resp_img.read().decode())
    print(f"POST /query (image) => {resp_img.status}")
    print(json.dumps(query_data_img, indent=2))
    conn.close()
    
    print("\n--- Testing Query Endpoint (All) ---")
    query_body_all = json.dumps({
        "kb_id": test_kb_id,
        "query": "search query for both text and images about architecture",
        "top_k": 5,
        "modality_filter": "all"
    })
    conn = http.client.HTTPConnection('127.0.0.1', 8000)
    conn.request('POST', '/api/v1/query/', body=query_body_all, headers={'Content-Type': 'application/json'})
    resp_all = conn.getresponse()
    query_data_all = json.loads(resp_all.read().decode())
    print(f"POST /query (all) => {resp_all.status}")
    print(json.dumps(query_data_all, indent=2))
    conn.close()
    
    print("\n--- Testing Delete Endpoint ---")
    if doc_id_image:
        conn = http.client.HTTPConnection('127.0.0.1', 8000)
        conn.request('DELETE', f'/api/v1/ingest/{doc_id_image}')
        resp5 = conn.getresponse()
        del_data5 = json.loads(resp5.read().decode())
        print(f"DELETE Image /{doc_id_image} => {resp5.status}")
        print(json.dumps(del_data5, indent=2))
        conn.close()
        
    db.close()
